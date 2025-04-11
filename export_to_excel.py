from docx import Document
import pandas as pd

def extract_tables_from_word(word_doc_path):
    try:
        # Load the Word document
        doc = Document(word_doc_path)
        
        all_tables_data = []
        
        # Iterate through each table in the document
        for i, table in enumerate(doc.tables, 1):
            table_data = []
            # Iterate through each row in the table
            for row in table.rows:
                row_data = []
                # Iterate through each cell in the row
                for cell in row.cells:
                    # Extract text from the cell
                    cell_text = ' '.join(paragraph.text for paragraph in cell.paragraphs).strip()
                    row_data.append(cell_text)
                
                # Only add the row if it doesn't contain "Insert your response here ..."
                if not any("Insert your response here" in cell for cell in row_data):
                    table_data.append(row_data)
            
            all_tables_data.append(table_data)
        
        return all_tables_data
    
    except Exception as e:
        print(f"An error occurred while reading the Word document: {str(e)}")
        return None

def combine_tables_to_single_dataframe(tables_data):
    if not tables_data or len(tables_data) == 0:
        print("No tables found or error in extraction.")
        return None
    
    all_dataframes = []
    
    # Process each table
    for i, table in enumerate(tables_data, 1):
        if not table or len(table) < 1:  # Skip empty tables
            continue
            
        # Use the first row as headers if available
        headers = table[0] if table[0] else [f'Column_{j}' for j in range(len(table[0]))]
        
        # Make headers unique by adding a suffix if needed
        unique_headers = []
        seen_headers = {}
        for header in headers:
            if header in seen_headers:
                seen_headers[header] += 1
                unique_headers.append(f"{header}_{seen_headers[header]}")
            else:
                seen_headers[header] = 0
                unique_headers.append(header)
        
        # Convert remaining rows to data (skip header row)
        data_rows = table[1:] if len(table) > 1 else []
        
        # Create a DataFrame for this table
        df = pd.DataFrame(data_rows, columns=unique_headers)
        
        # Add a column to indicate which table the data came from
        df['Table_Source'] = f'Table_{i}'
        
        all_dataframes.append(df)
    
    # Combine all DataFrames into one
    if not all_dataframes:
        return None
    
    combined_df = pd.concat(all_dataframes, ignore_index=True)
    
    return combined_df

def save_to_excel(df, output_excel_path):
    try:
        if df is not None and not df.empty:
            # Save to Excel
            df.to_excel(output_excel_path, sheet_name='Combined_Tables', index=False)
            print(f"Combined table saved to Excel file: {output_excel_path}")
        else:
            print("No data to save to Excel.")
            
    except Exception as e:
        print(f"An error occurred while writing to Excel: {str(e)}")

# Example usage
if __name__ == "__main__":
    # Specify the path to your Word document and the output Excel file
    word_doc_path = "./questions.docx"  # Replace with your .docx file path
    output_excel_path = "combined_tables.xlsx"  # Output Excel file path
    
    # Extract tables from the Word document
    tables_data = extract_tables_from_word(word_doc_path)
    
    if tables_data:
        # Combine all tables into a single DataFrame
        combined_df = combine_tables_to_single_dataframe(tables_data)
        
        if combined_df is not None:
            # Save to Excel
            save_to_excel(combined_df, output_excel_path)
        else:
            print("Failed to combine tables.")
    else:
        print("No tables were extracted or an error occurred.")
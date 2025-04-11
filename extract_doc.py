from docx import Document

def extract_table_text(word_doc_path):
    try:
        # Load the Word document
        doc = Document(word_doc_path)
        
        # Initialize a list to store all table text
        all_table_text = []
        
        # Iterate through each table in the document
        for table in doc.tables:
            table_text = []
            # Iterate through each row in the table
            for row in table.rows:
                row_text = []
                # Iterate through each cell in the row
                for cell in row.cells:
                    # Extract text from the cell (including paragraphs)
                    cell_text = ''
                    for paragraph in cell.paragraphs:
                        cell_text += paragraph.text + ' '
                    row_text.append(cell_text.strip())
                table_text.append(row_text)
            
            all_table_text.append(table_text)
        
        return all_table_text
    
    except Exception as e:
        print(f"An error occurred: {str(e)}")
        return None

def print_table_text(table_data):
    if table_data is None or len(table_data) == 0:
        print("No tables found or error in extraction.")
        return
    
    for i, table in enumerate(table_data, 1):
        print(f"\nTable {i}:")
        print("-" * 50)
        for row in table:
            print("| " + " | ".join(row) + " |")
        print("-" * 50)

# Example usage
if __name__ == "__main__":
    # Specify the path to your Word document
    word_doc_path = "your_document.docx"  # Replace with your .docx file path
    
    # Extract text from tables
    table_data = extract_table_text(word_doc_path)
    
    # Print the extracted text
    print_table_text(table_data)
    
    # If you want to save the text to a file or process it further, you can do so here
    if table_data:
        with open("table_output.txt", "w", encoding="utf-8") as f:
            for i, table in enumerate(table_data, 1):
                f.write(f"Table {i}:\n")
                f.write("-" * 50 + "\n")
                for row in table:
                    f.write("| " + " | ".join(row) + " |\n")
                f.write("-" * 50 + "\n")
        print("Table text has been saved to 'table_output.txt'")